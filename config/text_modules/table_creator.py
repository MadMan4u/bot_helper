from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_table(data, output_filename_word):
    unique_phrases = data

    # Создание документа Word
    doc = Document()

    # Установка альбомной ориентации
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)  # Ширина A4 в альбомной ориентации
    section.page_height = Inches(8.27)  # Высота A4 в альбомной ориентации

    # Установка шрифта по умолчанию (Times New Roman, 11)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Добавление таблицы
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False  # Отключить автонастройку ширины столбцов

    # Задание новой ширины столбцов
    table.columns[0].width = Inches(21 / 72)  # Столбец № (21px)
    table.columns[1].width = Inches(55 / 72)  # Столбец Намерение (55px)
    table.columns[2].width = Inches(248 / 72)  # Столбец Речевой модуль (248px)

    # Заголовки таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Намерение'
    hdr_cells[2].text = 'Речевой модуль'

    # Центрирование текста в заголовках и установка жирного шрифта
    for cell in hdr_cells:
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.bold = True
        paragraph.alignment = 1  # Центрирование
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Добавление обводки таблицы
    tbl = table._tbl  # Получаем таблицу как XML
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tbl.tblPr.append(tblBorders)

    # Заполнение таблицы уникальными речевыми модулями
    for i, (phrase, intent_text) in enumerate(unique_phrases, 1):
        row_cells = table.add_row().cells

        # Вставляем динамическую нумерацию с использованием стиля списка Word
        paragraph = row_cells[0].paragraphs[0]
        run = paragraph.add_run()
        run.add_break()
        paragraph.style = 'List Number'
        run.bold = True

        row_cells[1].text = intent_text if intent_text else ''
        row_cells[2].text = phrase

        # Форматирование текста
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

    # Сохранение документа Word
    doc.save(output_filename_word)
